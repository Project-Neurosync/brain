"""
NeuroSync RAG Demo Application

This standalone demo showcases the core document ingestion, storage, and retrieval
capabilities of NeuroSync using Llama and ChromaDB.
"""

import os
import streamlit as st
import tempfile
from typing import List, Dict, Any
import time
import uuid

# Imports for llama-index 0.9.48 (using old-style paths)
from llama_index import SimpleDirectoryReader, VectorStoreIndex, StorageContext
from llama_index.llms.ollama import Ollama
from llama_index.vector_stores.chroma import ChromaVectorStore
from llama_index.embeddings.huggingface import HuggingFaceEmbedding
from llama_index.node_parser import SimpleNodeParser
from llama_index.schema import Document
from llama_index.indices.service_context import ServiceContext
from llama_index import set_global_service_context
from llama_index.response_synthesizers import get_response_synthesizer
from llama_index.prompts.prompts import QuestionAnswerPrompt

import chromadb
import fitz  # PyMuPDF
import docx
import pandas as pd

# Configure page
st.set_page_config(
    page_title="NeuroSync RAG Demo",
    page_icon="🧠",
    layout="wide"
)

# Constants
CHROMA_PERSIST_DIR = os.environ.get("CHROMA_PERSIST_DIR", "chroma_db")

# Initialize session state
if "documents" not in st.session_state:
    st.session_state.documents = []
if "processed_files" not in st.session_state:
    st.session_state.processed_files = set()
if "index" not in st.session_state:
    st.session_state.index = None
if "query_history" not in st.session_state:
    st.session_state.query_history = []

def initialize_llm():
    """Initialize the Llama model using Ollama."""
    
    # Use Ollama API to access the llama3 model
    llm = Ollama(
        model="llama3:8b-instruct-q4_K_M",
        temperature=0.1,
        request_timeout=120.0,  # Increase timeout to 120 seconds
    )
    
    return llm

def initialize_chroma_db():
    """Initialize ChromaDB."""
    os.makedirs(CHROMA_PERSIST_DIR, exist_ok=True)
    chroma_client = chromadb.PersistentClient(path=CHROMA_PERSIST_DIR)
    collection_name = "neurosync_demo"
    
    # Get or create collection
    try:
        collection = chroma_client.get_collection(collection_name)
    except:
        collection = chroma_client.create_collection(collection_name)
    
    vector_store = ChromaVectorStore(chroma_collection=collection)
    storage_context = StorageContext.from_defaults(vector_store=vector_store)
    
    return storage_context

def extract_text_from_pdf(file_path):
    """Extract text from PDF files."""
    doc = fitz.open(file_path)
    text = ""
    for page in doc:
        text += page.get_text()
    return text

def extract_text_from_docx(file_path):
    """Extract text from DOCX files."""
    doc = docx.Document(file_path)
    text = ""
    for para in doc.paragraphs:
        text += para.text + "\n"
    return text

def process_file(file, file_key):
    """Process a file and extract its contents."""
    if file_key in st.session_state.processed_files:
        return None
    
    with tempfile.NamedTemporaryFile(delete=False, suffix=f".{file.name.split('.')[-1]}") as tmp:
        tmp.write(file.getvalue())
        tmp_path = tmp.name
    
    try:
        if file.name.lower().endswith('.pdf'):
            text = extract_text_from_pdf(tmp_path)
        elif file.name.lower().endswith('.docx'):
            text = extract_text_from_docx(tmp_path)
        elif file.name.lower().endswith(('.txt', '.md')):
            text = file.getvalue().decode('utf-8')
        elif file.name.lower().endswith('.csv'):
            df = pd.read_csv(tmp_path)
            text = df.to_markdown()
        else:
            st.error(f"Unsupported file format: {file.name}")
            return None
        
        doc_id = str(uuid.uuid4())
        doc = Document(text=text, metadata={"filename": file.name, "id": doc_id})
        st.session_state.documents.append(doc)
        st.session_state.processed_files.add(file_key)
        
        return doc
    finally:
        os.unlink(tmp_path)

def build_index():
    """Build the vector index from documents."""
    if not st.session_state.documents:
        st.warning("Please upload documents first!")
        return
    
    with st.spinner("Building knowledge index... (this might take a while)"):
        # Initialize embedding model
        embed_model = HuggingFaceEmbedding(model_name="sentence-transformers/all-MiniLM-L6-v2")
        
        # Initialize LLM
        llm = initialize_llm()
        
        # Set up node parser for chunking
        node_parser = SimpleNodeParser.from_defaults(
            chunk_size=1024,
            chunk_overlap=100
        )
        
        # Set up service context
        service_context = ServiceContext.from_defaults(
            llm=llm, 
            embed_model=embed_model,
            node_parser=node_parser
        )
        
        set_global_service_context(service_context)
        
        # Set up storage context with ChromaDB
        storage_context = initialize_chroma_db()
        
        # Build index
        st.session_state.index = VectorStoreIndex.from_documents(
            st.session_state.documents,
            storage_context=storage_context,
            service_context=service_context,
            show_progress=True
        )
        
        st.success(f"Knowledge base created with {len(st.session_state.documents)} documents!")

def answer_query(query):
    """Query the index and get an answer."""
    if not st.session_state.index:
        st.warning("Please build the knowledge index first!")
        return
    
    with st.spinner("Thinking..."):
        # Confiaccorgure query engine for more detailed responses
        response_synthesizer = get_response_synthesizer(
            response_mode="tree_summarize",
            streaming=False,
            verbose=True,
            structured_answer_filtering=False
        )
        
        # Customize the text QA template for more detailed responses
        qa_template = QuestionAnswerPrompt(
            """You are an expert AI research assistant providing comprehensive and detailed information based on technical documents. 
            Your task is to answer the question thoroughly and in detail, using all the relevant context provided.
            
            The context information is below.
            ---------------------
            {context_str}
            ---------------------
            
            Given the context information and not prior knowledge, provide a detailed, comprehensive, and well-structured answer to the question: {query_str}
            
            Your answer should:
            1. Be exhaustive and cover all aspects mentioned in the context
            2. Include specific technical details, code examples, and explanations where available
            3. Organize information in a clear structure with sections if appropriate
            4. Cite specific parts of the documents where relevant
            5. Be at least 4-5 paragraphs long unless the context is very limited
            
            Answer:"""
        )
        
        query_engine = st.session_state.index.as_query_engine(
            response_synthesizer=response_synthesizer,
            text_qa_template=qa_template,
            similarity_top_k=5,  # Increased from 3 to 5 for more context
            verbose=True,
        )
        
        response = query_engine.query(query)
        
        # Record history
        st.session_state.query_history.append({
            "query": query,
            "response": str(response),
            "timestamp": time.strftime("%Y-%m-%d %H:%M:%S")
        })
        
        return response

# Main UI Layout
st.title("🧠 NeuroSync RAG Demo")
st.subheader("Retrieval Augmented Generation with Llama")

st.markdown("""
This demo showcases the core RAG capabilities of NeuroSync:
1. Upload documents (PDF, DOCX, TXT, MD, CSV)
2. Process and index them with ChromaDB
3. Ask questions about your documents
4. Get accurate answers with relevant citations
""")

# Sidebar
st.sidebar.title("NeuroSync RAG Demo")

# Document Upload Section
st.sidebar.header("1. Upload Documents")
uploaded_files = st.sidebar.file_uploader(
    "Upload documents to your knowledge base", 
    accept_multiple_files=True,
    type=["pdf", "docx", "txt", "md", "csv"]
)

if uploaded_files:
    for file in uploaded_files:
        file_key = f"{file.name}_{file.size}"
        if file_key not in st.session_state.processed_files:
            if process_file(file, file_key):
                st.sidebar.success(f"Processed: {file.name}")

# Index Building Section
st.sidebar.header("2. Build Knowledge Index")
if st.sidebar.button("Build Index"):
    build_index()

# Document Stats
doc_count = len(st.session_state.documents)
st.sidebar.metric("Documents Processed", doc_count)

# Query Section
st.header("Ask Questions About Your Documents")
query = st.text_input("Enter your question:")

if st.button("Ask"):
    if query:
        response = answer_query(query)
        
        if response:
            st.markdown("### Answer")
            st.markdown(str(response))
            
            st.markdown("### Sources")
            for source_node in response.source_nodes:
                with st.expander(f"Source: {source_node.metadata.get('filename', 'Unknown')}"):
                    st.markdown(source_node.node.get_content())

# Query History Section
if st.session_state.query_history:
    st.header("Query History")
    for i, item in enumerate(reversed(st.session_state.query_history)):
        with st.expander(f"Q: {item['query']} ({item['timestamp']})"):
            st.markdown(item['response'])

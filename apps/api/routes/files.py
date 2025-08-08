"""
File upload and processing API routes for NeuroSync
Handles document upload, processing, and integration with RAG system
"""

import os
import uuid
import tempfile
from typing import List, Optional
from fastapi import APIRouter, Depends, HTTPException, UploadFile, File, Form, status
from fastapi.responses import JSONResponse
from sqlalchemy.orm import Session

from models.database import get_db
from models.user_models import User
from models.project import Project
from models.document import Document
from middleware.auth import get_current_user
from services.rag.rag_service import RAGService
from services.vectorstore.chroma import ChromaVectorStore
import logging

# File processing imports (matching demo)
import fitz  # PyMuPDF
import docx
import pandas as pd

logger = logging.getLogger(__name__)

# Create router
router = APIRouter(prefix="/api/v1/files", tags=["files"])

# Supported file types
SUPPORTED_EXTENSIONS = {'.txt', '.md', '.pdf', '.docx', '.doc', '.csv'}
MAX_FILE_SIZE = 10 * 1024 * 1024  # 10MB

def extract_text_from_pdf(file_path: str) -> str:
    """Extract text from PDF files using PyMuPDF."""
    text = ""
    try:
        doc = fitz.open(file_path)
        for page in doc:
            text += page.get_text()
        doc.close()
    except Exception as e:
        logger.error(f"Error extracting text from PDF: {e}")
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail=f"Failed to extract text from PDF: {str(e)}"
        )
    return text

def extract_text_from_docx(file_path: str) -> str:
    """Extract text from DOCX files."""
    try:
        doc = docx.Document(file_path)
        text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
    except Exception as e:
        logger.error(f"Error extracting text from DOCX: {e}")
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail=f"Failed to extract text from DOCX: {str(e)}"
        )
    return text

def process_file_content(file_content: bytes, filename: str) -> str:
    """Process file content and extract text based on file type."""
    file_ext = os.path.splitext(filename)[1].lower()
    
    # Create temporary file for processing
    with tempfile.NamedTemporaryFile(delete=False, suffix=file_ext) as tmp:
        tmp.write(file_content)
        tmp_path = tmp.name
    
    try:
        if file_ext == '.pdf':
            text = extract_text_from_pdf(tmp_path)
        elif file_ext in ['.docx', '.doc']:
            text = extract_text_from_docx(tmp_path)
        elif file_ext in ['.txt', '.md']:
            text = file_content.decode('utf-8')
        elif file_ext == '.csv':
            df = pd.read_csv(tmp_path)
            text = df.to_markdown()
        else:
            raise HTTPException(
                status_code=status.HTTP_400_BAD_REQUEST,
                detail=f"Unsupported file format: {file_ext}"
            )
        
        return text
    finally:
        # Clean up temporary file
        try:
            os.unlink(tmp_path)
        except:
            pass

@router.post("/upload")
async def upload_files(
    files: List[UploadFile] = File(...),
    project_id: Optional[str] = Form(None),
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    """
    Upload multiple files and process them for RAG
    Creates a new project if project_id is not provided
    """
    try:
        logger.info(f"File upload request from user: {current_user.email}")
        logger.info(f"Number of files: {len(files)}")
        logger.info(f"Project ID: {project_id}")
        
        # Validate files and extract content
        processed_files = []
        for file in files:
            # Check file extension
            file_ext = os.path.splitext(file.filename)[1].lower()
            if file_ext not in SUPPORTED_EXTENSIONS:
                raise HTTPException(
                    status_code=status.HTTP_400_BAD_REQUEST,
                    detail=f"Unsupported file type: {file_ext}. Supported types: {', '.join(SUPPORTED_EXTENSIONS)}"
                )
            
            # Check file size
            content = await file.read()
            if len(content) > MAX_FILE_SIZE:
                raise HTTPException(
                    status_code=status.HTTP_400_BAD_REQUEST,
                    detail=f"File {file.filename} is too large. Maximum size: 10MB"
                )
            
            # Extract text content from file
            logger.info(f"Processing file: {file.filename}")
            text_content = process_file_content(content, file.filename)
            
            processed_files.append({
                'filename': file.filename,
                'content': content,
                'text_content': text_content,
                'size': len(content),
                'type': file_ext
            })
            
            # Reset file pointer for potential reuse
            await file.seek(0)
        
        # Create project if not provided
        if not project_id:
            logger.info("Creating new project for file upload")
            new_project = Project(
                name="RAG Test Project",
                description="Project created for testing RAG with uploaded documents",
                status="active",
                progress=0,
                user_id=current_user.id
            )
            db.add(new_project)
            db.commit()
            db.refresh(new_project)
            project_id = str(new_project.id)
            logger.info(f"Created project: {project_id}")
        else:
            # Verify project exists and user has access
            project = db.query(Project).filter(
                Project.id == project_id,
                Project.user_id == current_user.id
            ).first()
            
            if not project:
                raise HTTPException(
                    status_code=status.HTTP_404_NOT_FOUND,
                    detail="Project not found or access denied"
                )
        
        # Process files and add to RAG system
        logger.info(f"Adding {len(processed_files)} files to RAG system for project {project_id}")
        
        # Initialize RAG service
        vector_store = ChromaVectorStore()
        rag_service = RAGService(vector_store=vector_store)
        
        results = []
        documents_added = []
        
        for file_info in processed_files:
            try:
                # Create document record in database
                doc_id = str(uuid.uuid4())
                document = Document(
                    id=doc_id,
                    title=file_info['filename'],
                    content=file_info['text_content'],
                    source="file_upload",
                    metadata={
                        "filename": file_info['filename'],
                        "file_type": file_info['type'],
                        "file_size": file_info['size'],
                        "project_id": project_id
                    },
                    project_id=project_id,
                    user_id=current_user.id
                )
                
                db.add(document)
                documents_added.append(document)
                
                # Add to vector store for RAG
                await rag_service.add_document(
                    doc_id=doc_id,
                    content=file_info['text_content'],
                    metadata={
                        "title": file_info['filename'],
                        "source": "file_upload",
                        "project_id": project_id,
                        "filename": file_info['filename']
                    }
                )
                
                results.append({
                    'filename': file_info['filename'],
                    'size': file_info['size'],
                    'type': file_info['type'],
                    'status': 'processed',
                    'document_id': doc_id,
                    'message': 'File uploaded, processed, and indexed successfully'
                })
                
                logger.info(f"Successfully processed and indexed: {file_info['filename']}")
                
            except Exception as e:
                logger.error(f"Error processing file {file_info['filename']}: {str(e)}")
                results.append({
                    'filename': file_info['filename'],
                    'size': file_info['size'],
                    'type': file_info['type'],
                    'status': 'error',
                    'message': f'Failed to process file: {str(e)}'
                })
        
        # Commit all database changes
        db.commit()
        
        logger.info(f"Successfully processed {len([r for r in results if r['status'] == 'processed'])} files")
        
        return JSONResponse(
            status_code=status.HTTP_200_OK,
            content={
                'message': f'Successfully uploaded and processed {len([r for r in results if r["status"] == "processed"])} files',
                'project_id': project_id,
                'files': results,
                'total_files': len(results),
                'total_size': sum(f['size'] for f in results),
                'processed_count': len([r for r in results if r['status'] == 'processed']),
                'error_count': len([r for r in results if r['status'] == 'error'])
            }
        )
        
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error uploading files: {str(e)}")
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=f"Failed to upload files: {str(e)}"
        )

@router.get("/supported-types")
async def get_supported_file_types():
    """
    Get list of supported file types for upload
    """
    return {
        'supported_extensions': list(SUPPORTED_EXTENSIONS),
        'max_file_size_mb': MAX_FILE_SIZE // (1024 * 1024),
        'description': 'Supported file types for document upload and RAG processing'
    }

@router.get("/project/{project_id}/files")
async def get_project_files(
    project_id: str,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    """
    Get list of files uploaded to a specific project
    """
    try:
        # Verify project exists and user has access
        project = db.query(Project).filter(
            Project.id == project_id,
            Project.user_id == current_user.id
        ).first()
        
        if not project:
            raise HTTPException(
                status_code=status.HTTP_404_NOT_FOUND,
                detail="Project not found or access denied"
            )
        
        # In a full implementation, this would query the documents table
        # For now, return a placeholder response
        return {
            'project_id': project_id,
            'project_name': project.name,
            'files': [],
            'total_files': 0,
            'message': 'File listing not yet implemented - files are processed and stored in vector database'
        }
        
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error getting project files: {str(e)}")
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=f"Failed to get project files: {str(e)}"
        )
